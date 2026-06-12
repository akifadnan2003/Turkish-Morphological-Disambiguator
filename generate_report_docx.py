"""
Generate Word (.docx) report for Turkish Morphological Disambiguation.
Run: python -X utf8 generate_report_docx.py
Output: report/NLP_Project_Report_AkifAdnan.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = os.path.dirname(os.path.abspath(__file__))
RES     = os.path.join(BASE, "results")
OUT_DIR = os.path.join(BASE, "report")
os.makedirs(OUT_DIR, exist_ok=True)
OUT_DOCX = os.path.join(OUT_DIR, "NLP_Project_Report_AkifAdnan.docx")

STUDENT_NAME  = "Akif Adnan"
STUDENT_NO    = "20360859106"
COURSE        = "Doğal Dil İşleme (Natural Language Processing)"
SEMESTER      = "2025 – 2026 Bahar Dönemi"
UNIVERSITY    = "T.C. Bursa Teknik Üniversitesi"
FACULTY       = "Mühendislik ve Doğa Bilimleri Fakültesi"
DEPARTMENT    = "Bilgisayar Mühendisliği Bölümü"

NAVY  = RGBColor(0x1d, 0x4e, 0xd8)
DKBLUE= RGBColor(0x1e, 0x3a, 0x5f)
GRAY  = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x15, 0x80, 0x3d)
AMBER = RGBColor(0x92, 0x40, 0x0e)
RED   = RGBColor(0x99, 0x1b, 0x1b)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xf1, 0xf5, 0xf9)
NAVY_FILL = "1d4ed8"
LGRAY_FILL = "f1f5f9"
BLUE_FILL  = "dbeafe"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = NAVY if level == 1 else DKBLUE
    run.font.name = "Calibri"
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def make_header_row(table, headers, col_widths_cm=None):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        set_cell_bg(cell, NAVY_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for r in table.rows:
                r.cells[i].width = Cm(w)


def add_data_row(table, values, center_cols=None, bold=False, fill=None):
    center_cols = center_cols or []
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].alignment = align
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.bold = bold
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if fill:
            set_cell_bg(cell, fill)
    return row


def f1_color(f1):
    if f1 >= 90: return GREEN
    if f1 >= 80: return AMBER
    return RED


# ── Build document ──────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── COVER ────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.add_run("\n\n")

for line in [UNIVERSITY, FACULTY, DEPARTMENT]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DKBLUE
    r.font.name = "Calibri"

doc.add_paragraph()

for line in [COURSE, SEMESTER]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    r.font.name = "Calibri"

doc.add_paragraph("\n")

p = doc.add_paragraph("Proje 1 — Morfolojik Çözümleme")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.runs[0]; r.font.size = Pt(12); r.font.color.rgb = GRAY; r.font.name = "Calibri"

p = doc.add_paragraph("Turkish Morphological Disambiguation")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.runs[0]; r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = NAVY; r.font.name = "Calibri"

doc.add_paragraph("\n\n")

for line in [STUDENT_NAME, f"Öğrenci No: {STUDENT_NO}", "Bireysel Çalışma", "Haziran 2026"]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]; r.font.size = Pt(11); r.font.name = "Calibri"

doc.add_page_break()

# ── 1. GİRİŞ ────────────────────────────────────────────────
heading(doc, "1. Giriş")
body(doc,
    "Morfolojik çözümleme, bir cümledeki her kelimenin birden fazla olası dilbilgisel çözümü "
    "bulunduğunda doğru çözümün seçilmesi görevidir. Bu proje, eklemeli bir dil olan Türkçe "
    "için bu problemi ele almaktadır. Türkçede kelimeler köklere son ekler eklenerek "
    "oluşturulduğundan, aynı yüzey biçimi farklı morfolojik okumalara sahip olabilmektedir."
)
doc.add_paragraph()
body(doc,
    "Proje hedefi: UD Türkçe-IMST Ağaç Bankası üzerinde eğitilmiş bir İstatistiksel Makine "
    "Öğrenmesi modeli (Koşullu Rasgele Alanlar — CRF) ile verilen bir Türkçe cümledeki "
    "her belirteç için doğru Evrensel Sözcük Türü (UPOS) etiketini tahmin etmek. Model; "
    "ayrılmış test kümesi ve 25 elle etiketlenmiş cümle üzerinde değerlendirilmiştir."
)

# ── 2. VERİ SETİ ────────────────────────────────────────────
heading(doc, "2. Veri Seti")
body(doc,
    "UD Türkçe-IMST Ağaç Bankası (Universal Dependencies projesi), İstanbul Teknik "
    "Üniversitesi tarafından oluşturulmuş olup CoNLL-U formatında dilbilimciler tarafından "
    "etiketlenmiş cümleler içermektedir."
)
doc.add_paragraph()

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
make_header_row(t, ["Bölünme", "Cümle", "Belirteç", "Amaç"],
                col_widths_cm=[3.5, 3, 3.5, 6])
for row_data in [
    ("Eğitim (Train)", "3.435",  "~52.000", "Model eğitimi"),
    ("Doğrulama (Dev)", "1.100", "~17.000", "Hiperparametre ayarı"),
    ("Test",           "1.100",  "~10.000", "Nihai değerlendirme"),
    ("Elle etiket (özel)", "25", "169",     "Öğrenci tarafından elle etiketlenmiş"),
]:
    add_data_row(t, row_data, center_cols=[1, 2])

doc.add_paragraph("Tablo 1. Projede kullanılan veri seti bölünmeleri.", style="Caption")

doc.add_paragraph()
body(doc,
    "CoNLL-U Formatı: Her cümle sekme ayrımlı bir dosyada saklanmaktadır. Temel sütunlar: "
    "ID (konum), FORM (yüzey kelime), LEMMA (kök), UPOS (tahmin edilen etiket) ve FEATS "
    "(Durum, Sayı, Zaman, Kişi gibi morfolojik özellikler)."
)

doc.add_page_break()

# ── 3. YÖNTEMBİLİM ──────────────────────────────────────────
heading(doc, "3. Yöntem")
heading(doc, "3.1 Algoritma: Koşullu Rasgele Alanlar (CRF)", level=2)
body(doc,
    "Koşullu Rasgele Alan, ayrımcı bir dizi etiketleme modelidir. Basit bir belirteç bazlı "
    "sınıflandırıcının aksine, CRF tüm cümleyi birlikte etiketler; komşu etiketler arasındaki "
    "geçiş olasılıklarını öğrenir (örneğin DET'den sonra neredeyse her zaman NOUN veya ADJ "
    "gelir). Bu özellik, sözcük sırasının güçlü bağlamsal ipuçları taşıdığı Türkçe için "
    "kritik öneme sahiptir. Model; c1=0.05 ve c2=0.05 düzenlileştirme değerleriyle "
    "L-BFGS optimizasyonu kullanılarak eğitilmiştir."
)

heading(doc, "3.2 Özellik Mühendisliği", level=2)
body(doc,
    "Türkçe morfolojisine özgü aşağıdaki özellikler el ile tasarlanmıştır:"
)
doc.add_paragraph()

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Table Grid"
make_header_row(t2, ["Özellik", "Örnekler", "Sinyal"],
                col_widths_cm=[5, 4.5, 5])
for row_data in [
    ("Son ekler (2–4 karakter)", "-dan/-den, -da/-de", "Ayrılma/Bulunma → NOUN"),
    ("Zaman eki", "-yor/-iyor, -dı/-di", "Geniş/Geçmiş Zaman → VERB"),
    ("Çoğul eki", "-lar/-ler", "NOUN"),
    ("Mastar eki", "-mak/-mek", "VERB (sözel ad)"),
    ("Sıfat yapım eki", "-lı/-lu, -sız/-suz", "ADJ"),
    ("Rivayet eki", "-mış/-müş", "VERB"),
    ("Ünlü uyumu", "Son ünlünün ön/arka sınıfı", "Morfolojik sınıf"),
    ("±2 belirteç bağlam penceresi", "Önceki/sonraki 2 kelime", "Bağlam ile çözümleme"),
    ("Büyük harf ile başlama", "İlk harf büyük mü?", "PROPN sinyali"),
    ("Kesme işareti var mı?", "Ankara'da", "PROPN (özel ad + ek)"),
    ("Rakam / rakam içeriyor", "2024, 3.5", "NUM"),
]:
    add_data_row(t2, row_data, center_cols=[])

doc.add_paragraph("Tablo 2. CRF modelinde kullanılan Türkçeye özgü özellikler (toplam 40+).", style="Caption")
doc.add_paragraph()

heading(doc, "3.3 Zemberek ve Aday Çözümleme Yaklaşımı (Candidate Generation)", level=2)
body(doc,
    "Proje isterlerinde belirtilen Zemberek tabanlı aday kelime çözümleme mantığı, "
    "sistemde Corpus-Based Candidate Dictionary (Derlem Tabanlı Aday Sözlüğü) "
    "yöntemiyle simüle edilmiştir. Java tabanlı Zemberek kütüphanesinin Python "
    "ortamında yaratabileceği bağımlılık ve gRPC köprüsü sorunlarını aşmak adına, "
    "eğitim veri setinde (UD_Turkish-IMST) geçen her kelimenin aldığı tüm olası "
    "morfolojik etiketler bir sözlükte toplanmıştır."
)
doc.add_paragraph()
body(doc,
    "candidate_disambig.py dosyasında görülebileceği üzere, sisteme verilen bir "
    "cümlenin her kelimesi için önce bu olası aday listesi (candidate list) "
    "çekilmekte, ardından CRF modelinin marjinal olasılıkları (predict_marginals) "
    "kullanılarak bu adaylar arasında en yüksek olasılıklı olanı seçilerek "
    "'Disambiguation' (Çözümleme) işlemi tamamlanmaktadır."
)

heading(doc, "3.4 İki Aşamalı Morfolojik Pipeline (Two-Stage Pipeline)", level=2)
body(doc, "Morfolojik özellik (FEATS) tahmini için iki aşamalı bir mimari benimsenmiştir:")
doc.add_paragraph()

p = doc.add_paragraph(style="List Bullet")
p.add_run("Aşama 1 — CRF: ").bold = True
p.add_run("Model, bağlam penceresindeki özellikler kullanılarak her belirteç için UPOS etiketini tahmin eder (%90,93 doğruluk).")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Aşama 2 — Derlem Sözlüğü: ").bold = True
p.add_run("Tahmin edilen (kelime, UPOS) çifti için eğitim verisindeki en sık görülen FEATS kombinasyonu atanır. Bu adım, Zemberek'in sağlayacağı morfolojik aday analizlerini simüle etmektedir.")

doc.add_paragraph()
body(doc,
    "Bu mimari tercih, 985 benzersiz UPOS+FEATS kombinasyonunu tek bir CRF modeliyle "
    "öğretmeye çalışmanın (sınıf patlaması — class explosion problemi) önüne geçerek "
    "akademik literatürde yaygın olan pipeline yaklaşımını uygulamaktadır."
)

doc.add_page_break()

# ── 4. ELLE ETİKETLEME ──────────────────────────────────────
heading(doc, "4. Elle Etiketleme")
body(doc,
    "Modeli bağımsız bir veri kümesi üzerinde değerlendirmek ve etiketleme görevinin "
    "anlaşıldığını göstermek amacıyla, CoNLL-U formatında 25 Türkçe cümle elle "
    "etiketlenmiştir. Bu cümleler ağaç bankasından alınmamış, farklı sözdizimsel "
    "yapıları ve alanları kapsayacak şekilde seçilmiştir."
)
doc.add_paragraph()

t3 = doc.add_table(rows=1, cols=3)
t3.style = "Table Grid"
make_header_row(t3, ["Alan", "Cümle Sayısı", "Örnek"],
                col_widths_cm=[4, 3, 8.5])
for row_data in [
    ("Üniversite / Akademik", "8", "Proje ödevini yarın teslim etmelisin ."),
    ("Günlük Yaşam", "8", "Kardeşim bu hafta İzmir'den gelecek ."),
    ("NLP / Teknoloji", "5", "CRF modeli Türkçe için başarılı sonuçlar verdi ."),
    ("Coğrafya", "2", "Türkiye'nin başkenti Ankara'dır ."),
    ("Doğa", "1", "Dağlarda kar yağıyordu ve hava çok soğuktu ."),
    ("NLP Değerlendirmesi", "1", "Modelin başarı oranı yüzde doksana ulaştı ."),
    ("Toplam", "25", "169 belirteç"),
]:
    add_data_row(t3, row_data, center_cols=[1])

doc.add_paragraph("Tablo 3. Elle etiketleme kümesi alanlara göre.", style="Caption")
doc.add_paragraph()
body(doc,
    "Her belirteç UPOS etiketi, LEMMA (sözlük kök biçimi) ve FEATS (Durum, Sayı, Kişi, "
    "Zaman, Kip, Kutupluluk, EylemBiçimi, Çatı, İyelik) ile etiketlenmiştir. "
    "Dönüşlü eylemler, sözel adlar, edilgen çatı ve olumsuzluk gibi dilbilimsel açıdan "
    "ilgi çekici yapılar da etiketlenmiştir."
)

doc.add_page_break()

# ── 5. SONUÇLAR ─────────────────────────────────────────────
heading(doc, "5. Sonuçlar")
heading(doc, "5.1 Test Kümesi (1.100 cümle, 10.032 belirteç)", level=2)
body(doc,
    "Eğitilen CRF modeli, UD Türkçe-IMST ağaç bankasının ayrılmış test bölümü üzerinde "
    "değerlendirilmiştir. Aşağıdaki tablo, her UPOS sınıfı için hassasiyet, duyarlılık, "
    "F1 puanı ve destek (örnek sayısı) değerlerini göstermektedir."
)
doc.add_paragraph()

TEST_METRICS = [
    ("PUNCT",  99.9,  100.0, 99.95, 1933),
    ("AUX",    94.4,  96.2,  95.3,  211),
    ("CCONJ",  96.4,  96.4,  96.4,  356),
    ("VERB",   93.8,  94.1,  93.9,  1928),
    ("PRON",   94.5,  92.9,  93.7,  464),
    ("ADP",    95.6,  90.5,  92.9,  357),
    ("DET",    84.2,  96.2,  89.8,  344),
    ("NOUN",   84.3,  92.5,  88.2,  2430),
    ("ADV",    88.8,  77.0,  82.5,  461),
    ("NUM",    90.7,  71.4,  79.9,  192),
    ("ADJ",    84.9,  75.8,  80.1,  960),
    ("PROPN",  83.3,  70.9,  76.6,  374),
]

t4 = doc.add_table(rows=1, cols=5)
t4.style = "Table Grid"
make_header_row(t4, ["UPOS Etiketi", "Hassasiyet", "Duyarlılık", "F1", "Destek"],
                col_widths_cm=[3.5, 3, 3, 2.5, 3])
for i, (tag, prec, rec, f1, sup) in enumerate(TEST_METRICS):
    fill = "f0fdf4" if f1 >= 90 else ("fffbeb" if f1 >= 80 else "fef2f2")
    row = add_data_row(t4,
        [tag, f"{prec:.1f}%", f"{rec:.1f}%", f"{f1:.2f}%", f"{sup:,}"],
        center_cols=[1, 2, 3, 4], fill=fill)
    run = row.cells[3].paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = f1_color(f1)

add_data_row(t4, ["Doğruluk", "", "", "90.93%", "10,032"],
             center_cols=[1, 2, 3, 4], bold=True, fill=BLUE_FILL)
add_data_row(t4, ["Ağırlıklı Ort.", "90.99%", "90.93%", "90.80%", "10,032"],
             center_cols=[1, 2, 3, 4], bold=True, fill=BLUE_FILL)

doc.add_paragraph("Tablo 4. Test kümesi üzerinde sınıf bazlı değerlendirme sonuçları.", style="Caption")

# Charts
heading(doc, "5.2 Sınıf Başına F1 Puanı", level=2)
body(doc,
    "Aşağıdaki grafikler test kümesindeki her UPOS sınıfı için F1 puanını ve "
    "karışıklık matrisini göstermektedir. Belirgin morfolojik işaretçileri olan sınıflar "
    "(PUNCT, AUX, VERB, CCONJ) neredeyse mükemmel puanlar elde etmiştir. En zor sınıflar "
    "PROPN ve ADJ olmuştur."
)

f1_img = os.path.join(RES, "f1_per_class_test.png")
if os.path.exists(f1_img):
    doc.add_picture(f1_img, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Şekil 1. Test kümesinde UPOS sınıfı başına F1 puanı.", style="Caption")

heading(doc, "5.3 Karışıklık Matrisi", level=2)
body(doc,
    "Karışıklık matrisi gerçek etiketleri (satırlar) tahmin edilen etiketlere (sütunlar) "
    "karşı göstermektedir. Köşegen doğru tahminleri temsil etmektedir. Temel hata kalıpları: "
    "PROPN sıklıkla NOUN olarak tahmin edilmekte; ADJ ise NOUN ile karıştırılmaktadır."
)

cm_img = os.path.join(RES, "confusion_matrix_test.png")
if os.path.exists(cm_img):
    doc.add_picture(cm_img, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Şekil 2. Test kümesi üzerinde karışıklık matrisi.", style="Caption")

doc.add_page_break()

# ── 6. ÖZEL DEĞERLENDİRME ───────────────────────────────────
heading(doc, "6. Elle Etiketleme Üzerinde Değerlendirme")
body(doc,
    "Eğitilen model, my_annotations.conllu dosyasındaki 25 elle etiketlenmiş cümle "
    "(169 belirteç) üzerinde de değerlendirilmiştir. Bu küme eğitim sırasında hiç "
    "görülmemiş, alan dışı bir değerlendirme imkânı sunmaktadır."
)
doc.add_paragraph()

CUSTOM_METRICS = [
    ("PUNCT",  100.0, 100.0, 100.0, 25),
    ("ADP",    100.0, 100.0, 100.0, 2),
    ("CCONJ",  100.0, 100.0, 100.0, 1),
    ("DET",    100.0, 100.0, 100.0, 11),
    ("PRON",   100.0, 100.0, 100.0, 2),
    ("VERB",   93.8,  100.0, 96.8,  30),
    ("NUM",    75.0,  100.0, 85.7,  3),
    ("PROPN",  75.0,  100.0, 85.7,  6),
    ("NOUN",   85.7,  82.8,  84.2,  58),
    ("ADJ",    63.6,  77.8,  70.0,  18),
    ("ADV",    100.0, 46.2,  63.2,  13),
]

t5 = doc.add_table(rows=1, cols=5)
t5.style = "Table Grid"
make_header_row(t5, ["UPOS Etiketi", "Hassasiyet", "Duyarlılık", "F1", "Destek"],
                col_widths_cm=[3.5, 3, 3, 2.5, 3])
for tag, prec, rec, f1, sup in CUSTOM_METRICS:
    fill = "f0fdf4" if f1 >= 90 else ("fffbeb" if f1 >= 80 else "fef2f2")
    row = add_data_row(t5,
        [tag, f"{prec:.1f}%", f"{rec:.1f}%", f"{f1:.1f}%", str(sup)],
        center_cols=[1, 2, 3, 4], fill=fill)
    run = row.cells[3].paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = f1_color(f1)

add_data_row(t5, ["Doğruluk", "", "", "87.57%", "169"],
             center_cols=[1, 2, 3, 4], bold=True, fill=BLUE_FILL)

doc.add_paragraph("Tablo 5. 25 elle etiketlenmiş cümle üzerinde sınıf bazlı değerlendirme.", style="Caption")
doc.add_paragraph()
body(doc,
    "Model özel kümede %87,57 doğruluk elde etmiş; bu sonuç modelin eğitim verisinin "
    "ötesine genelleşebildiğini doğrulamaktadır. Test kümesinden (%90,93) özel kümeye "
    "(%87,57) düşüş beklenen bir durumdur — özel cümleler daha karmaşık yapılar "
    "(dönüşlü eylemler, sözel adlar, edilgen çatı) ve alışılmışın dışında kelimeler "
    "içermektedir."
)

doc.add_page_break()

# ── 7. HATA ANALİZİ ─────────────────────────────────────────
heading(doc, "7. Hata Analizi")
body(doc,
    "Yanlış sınıflandırılan belirteçlerin analizi, üç sistematik hata kalıbını ortaya "
    "koymaktadır:"
)
heading(doc, "7.1 ADV ile NOUN Karışıklığı", level=2)
body(doc,
    "Dün, bugün, yarın gibi kısa zaman zarflarının belirgin bir eki bulunmamakta ve "
    "eğitim verisinde sıklıkla NOUN olarak geçmektedir. Bağlam yetersiz kaldığında model "
    "bu kelimeleri NOUN olarak etiketleme eğilimindedir. Bu durum, özel kümede ADV "
    "duyarlılığının %46'ya düşmesine yol açmaktadır."
)
heading(doc, "7.2 ADJ / NOUN Sınırı", level=2)
body(doc,
    "Türkçe sıfatlar ve isimler çoğunlukla aynı yüzey biçimini paylaşmaktadır. Model, "
    "izleyen bir isim veya belirteç olmadığında varsayılan olarak daha yaygın sınıf olan "
    "NOUN'u seçmektedir. Bu durum, her iki değerlendirme kümesindeki en büyük tek hata "
    "kaynağıdır."
)
heading(doc, "7.3 PROPN Hassasiyeti", level=2)
body(doc,
    "Ek alan özel adlar (Ankara'da, Türkiye'nin) kesme işareti özelliği sayesinde "
    "doğru tanınmaktadır. Ancak kesme işareti olmayan görülmemiş özel adlar sıklıkla "
    "NOUN olarak yanlış sınıflandırılmaktadır. Öte yandan model, özel ad olmayan büyük "
    "harfle başlayan cümle başı kelimeleri için PROPN'u fazla tahmin edebilmektedir."
)

# ── 8. SONUÇ ────────────────────────────────────────────────
heading(doc, "8. Sonuç")
body(doc,
    "Bu proje, Türkçe için Koşullu Rasgele Alanlar kullanılarak eksiksiz bir morfolojik "
    "çözümleme sistemi gerçekleştirmektedir. Model; standart UD Türkçe-IMST test kümesinde "
    "%90,93 doğruluk ve bağımsız olarak elle etiketlenmiş 25 cümlede %87,57 doğruluk "
    "elde etmiştir."
)
doc.add_paragraph()
body(doc,
    "Temel katkılar: (1) 40'tan fazla morfolojik özellik ile Türkçeye özgü özellik "
    "mühendisliği; (2) CoNLL-U formatında tam etiketleme hattı (eğitim verisi + elle "
    "etiketleme); (3) canlı gösterim için Flask web uygulaması; (4) sınıf bazlı "
    "değerlendirme (hassasiyet, duyarlılık, F1, doğruluk, karışıklık matrisi)."
)
doc.add_paragraph()
body(doc,
    "Temel sınırlama ADV/NOUN ve ADJ/NOUN karışıklığıdır. Gelecekteki çalışmalar, "
    "ek girdi özelliği olarak Zemberek gibi özel bir Türkçe analizörü veya daha uzun "
    "menzilli bağlamı yakalayan sinirsel bir dizi modeli (BiLSTM-CRF) kullanabilir."
)

# ── KAYDET ──────────────────────────────────────────────────
doc.save(OUT_DOCX)
print(f"\nWord raporu oluşturuldu: {OUT_DOCX}\n")
